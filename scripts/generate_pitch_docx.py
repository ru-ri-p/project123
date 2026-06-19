#!/usr/bin/env python3
"""Generate one-page Word brief from compliance content."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_sub(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_section(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(4)
    for r in p2.runs:
        r.font.size = Pt(8.5)
    p2.paragraph_format.space_before = Pt(0)


def _add_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(7.5)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, val in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main() -> None:
    out = ROOT / "docs" / "PITCH_COMPLIANCE_BRIEF.docx"
    doc = Document()
    _set_narrow_margins(doc)

    _add_heading(doc, "Attest — Regulatory Alignment Brief")
    _add_sub(doc, "UAE PDPL & CBUAE AI Guidance | June 2026 | One-page summary for LFIs & investors")

    _add_section(
        doc,
        "Executive summary",
        "Attest is a runtime governance and tamper-evident provenance control plane for AI workflows "
        "(UAE/MENA, financial sector first). It helps institutions document, enforce, and independently "
        "prove what AI did, when, under which policies — and that records were not altered. Attest supports "
        "PDPL-aligned technical controls and aligns with CBUAE AI supervisory expectations; it does not "
        "certify legal compliance or judge AI truthfulness. The institution’s board owns policy; Attest "
        "enforces and proves it was applied.",
    )

    _add_section(
        doc,
        "Legal context",
        "UAE has no standalone AI law — AI must comply with existing law (PDPL, sector rules, CBUAE "
        "supervision for banks). PDPL (Federal Decree-Law No. 45 of 2021) is binding (full compliance "
        "due 1 Jan 2027). CBUAE AI Guidance Note (11 Feb 2026) is non-binding but supervisory — frame as "
        "alignment with supervisory dialogue, not statute. DIFC/ADGM have separate regimes.",
    )

    p = doc.add_paragraph()
    r = p.add_run("PDPL — how Attest supports demonstrable controls")
    r.bold = True
    r.font.size = Pt(9)

    _add_table(
        doc,
        ("PDPL theme", "Attest capability"),
        [
            ("Minimisation & privacy by design", "PII redaction before storage"),
            ("Security of processing", "Encryption, signing, tenant isolation"),
            ("Right to erasure", "Crypto-shred + signed erasure proof"),
            ("Accountability / audit", "Offline verify.py on evidence export"),
            ("Processing records (ROPA-style)", "Traces = auditable AI workflow log"),
            ("DPIA readiness", "Structured evidence bundle for auditors"),
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("CBUAE AI Guidance (LFIs) — supervisory themes operationalised")
    r.bold = True
    r.font.size = Pt(9)

    _add_table(
        doc,
        ("CBUAE theme", "Attest capability"),
        [
            ("AI governance framework", "Versioned policies + policy_decision events"),
            ("Activity inventory", "Trace log per AI workflow"),
            ("Board accountability", "Named approver; approval_action events"),
            ("Human oversight", "RED/orange → approval queue; fail-closed default"),
            ("Explainability", "Reasons + policy version on every precheck"),
            ("Risk controls", "Tiers: green / yellow / orange / red"),
        ],
    )

    _add_section(
        doc,
        "60-second proof",
        "Precheck → tier + reasons → human approve/deny (if orange/red) → signed events → "
        "export evidence ZIP → auditor runs verify.py offline → ALL EVENTS VERIFIED. "
        "Erasure: destroy content, chain proves erasure occurred.",
    )

    _add_section(
        doc,
        "What we are / are not",
        "ARE: control + evidence layer; enforcer of your policies; tamper-evident trail. "
        "ARE NOT: legal advisor; compliance certificate; proof AI is correct. "
        "Status: Weeks 1–8 delivered. Ask: one CBUAE LFI design-partner pilot.",
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Disclaimer: Attest supports controls aligned with PDPL and CBUAE AI Guidance themes. "
        "Compliance is an institutional obligation — align with qualified UAE counsel. "
        "CBUAE Guidance is supervisory, not statute. Free zones may be outside mainland PDPL."
    )
    run.italic = True
    run.font.size = Pt(7)

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
